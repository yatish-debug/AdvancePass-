import os
import csv
import json
from typing import Optional, List, Dict, Any
from fpdf import FPDF, XPos, YPos
from docx import Document

REPORTS_DIR = os.path.join(os.path.dirname(__file__), "..", "data", "exports")
os.makedirs(REPORTS_DIR, exist_ok=True)

def _get_filename(base_name: str, ext: str) -> str:
    return os.path.join(REPORTS_DIR, f"{base_name}.{ext}")

def generate_csv_report(data: List[Dict[str, Any]], base_name: str = "security_report") -> str:
    """Generates a CSV report from a list of dictionaries."""
    if not data:
        return ""
    filepath = _get_filename(base_name, "csv")
    keys = list(data[0].keys())
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        writer.writerows(data)
    return filepath

def generate_json_report(data: List[Dict[str, Any]], base_name: str = "security_report") -> str:
    """Generates a JSON report."""
    filepath = _get_filename(base_name, "json")
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4)
    return filepath

def generate_pdf_report(data: List[Dict[str, Any]], base_name: str = "security_report") -> str:
    """Generates a professional PDF report."""
    filepath = _get_filename(base_name, "pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 16)
    pdf.cell(200, 10, text="AdvancePass V6.0 Security Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    pdf.set_font("helvetica", size=12)
    pdf.ln(10)

    for item in data:
        for k, v in item.items():
            pdf.cell(200, 8, text=f"{k}: {v}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(5)

    pdf.output(filepath)
    return filepath

def generate_docx_report(data: List[Dict[str, Any]], base_name: str = "security_report") -> str:
    """Generates a DOCX report."""
    filepath = _get_filename(base_name, "docx")
    doc = Document()
    doc.add_heading('AdvancePass V6.0 Security Report', 0)
    
    for item in data:
        p = doc.add_paragraph()
        for k, v in item.items():
            p.add_run(f"{k}: ").bold = True
            p.add_run(f"{v}\n")
            
    doc.save(filepath)
    return filepath

def generate_dashboard_pdf(metrics: Dict[str, Any], chart1_path: Optional[str] = None, chart2_path: Optional[str] = None, base_name: str = "dashboard_report") -> str:
    """Generates a PDF report for the Dashboard metrics with charts."""
    filepath = _get_filename(base_name, "pdf")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", 'B', 16)
    pdf.cell(200, 10, text="AdvancePass V6.0 Enterprise Dashboard Report", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

    pdf.set_font("helvetica", size=12)
    pdf.ln(10)

    pdf.cell(200, 10, text="Summary Metrics:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(5)

    def add_metric_row(label: str, value: Any):
        pdf.set_font("helvetica", 'B', 12)
        pdf.cell(80, 8, text=label, new_x=XPos.RIGHT, new_y=YPos.TOP)
        pdf.set_font("helvetica", size=12)
        pdf.cell(100, 8, text=str(value), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    total = int(metrics.get("total", 0))
    weak = int(metrics.get("weak_count", 0))
    strong = int(metrics.get("strong_count", 0))

    add_metric_row("Total Passwords Analyzed:", total)
    add_metric_row("Average Security Score:", f"{metrics.get('avg_score', 0)}/100")
    add_metric_row("Weak Passwords Detected:", weak)
    add_metric_row("Moderate Passwords:", total - weak - strong)
    add_metric_row("Strong Passwords:", strong)

    pdf.ln(10)
    
    # Embed charts
    y_pos = pdf.get_y()
    if chart1_path and os.path.exists(chart1_path):
        pdf.cell(200, 10, text="Strength Distribution", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.image(chart1_path, x=10, y=pdf.get_y(), w=90)

    if chart2_path and os.path.exists(chart2_path):
        pdf.set_y(y_pos)
        pdf.cell(200, 10, text="Entropy Distribution", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='R')
        pdf.image(chart2_path, x=110, y=pdf.get_y(), w=90)

    pdf.output(filepath)
    return filepath

